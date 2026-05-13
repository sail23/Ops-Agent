"""报告导出模块 - 支持 Word 和 PDF 格式"""
import io
import re
from datetime import datetime
from typing import Optional
from xml.sax.saxutils import escape as _xml_escape

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib import colors

from power_aiops.orchestration.debate import DebateResult


class ReportExporter:
    """辩论报告导出器"""

    def __init__(self, result: DebateResult):
        self.result = result

    def export_docx(self) -> bytes:
        """导出为 Word 文档"""
        doc = Document()

        # 设置文档标题
        title = doc.add_heading('智能运维故障分析报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 元信息
        doc.add_paragraph()
        meta_info = doc.add_paragraph()
        meta_info.add_run('故障编号：').bold = True
        meta_info.add_run(f"{self.result.incident_id or 'N/A'}    ")
        meta_info.add_run('生成时间：').bold = True
        meta_info.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))

        doc.add_paragraph()

        # 执行摘要
        doc.add_heading('执行摘要', level=1)
        summary = doc.add_paragraph()
        summary.add_run('辩论轮次：').bold = True
        summary.add_run(f"{self.result.total_turns or 0} 轮\n")
        summary.add_run('收敛分数：').bold = True
        summary.add_run(f"{((self.result.convergence_score or 0) * 100):.0f}%\n")
        summary.add_run('终止原因：').bold = True
        summary.add_run(f"{self.result.termination_reason or '正常结束'}\n")

        if self.result.disputed_points:
            summary.add_run('\n争议点：').bold = True
            for point in self.result.disputed_points:
                doc.add_paragraph(point, style='List Bullet')

        doc.add_paragraph()

        # 结论报告
        if self.result.report_text:
            doc.add_heading('分析结论', level=1)
            # 处理 Markdown 格式
            content = self.result.report_text
            for line in content.split('\n'):
                if line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('- '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line.strip():
                    doc.add_paragraph(line)

        doc.add_paragraph()

        # 辩论历史
        doc.add_heading('辩论过程', level=1)
        if self.result.history and 'turns' in self.result.history:
            for i, turn in enumerate(self.result.history['turns']):
                # 轮次标题
                p = doc.add_paragraph()
                p.add_run(f'第 {turn.get("turn_id", i + 1)} 轮 - {turn.get("round", "")}').bold = True

                # Agent 信息
                p2 = doc.add_paragraph()
                p2.add_run(f'发言人：{turn.get("agent_id", "N/A")}\n')

                # 内容
                if turn.get('message'):
                    msg = turn['message']
                    content = msg.get('content', '')
                    if content:
                        doc.add_paragraph(content[:500] + ('...' if len(content) > 500 else ''))

                doc.add_paragraph()  # 空行分隔

        # 页脚
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.add_run('— 由 AI-OPS 多智能体系统自动生成 —').italic = True
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 保存到字节流
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _escape_xml(self, text: str) -> str:
        """XML-escape text content, then convert Markdown inline formatting to reportlab XML tags.

        The order is critical: escape first so that stray < > & in the original
        text don't get mistaken for XML tags during the Markdown→XML step.
        Only **bold**, *italic*, and `code` are converted. Underscore-based
        formatting is deliberately NOT supported because `_` is ubiquitous in
        technical identifiers (Threads_connected, Max_connections, etc.).
        """
        if not text:
            return text

        text = _xml_escape(text)

        # Code (backticks)
        text = re.sub(r"`([^`]+)`", r"<font face='Courier'>\1</font>", text)
        # Bold: **text**
        text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
        # Italic: *text* (use lookarounds to avoid matching **bold**)
        text = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"<i>\1</i>", text)

        return text

    def export_pdf(self) -> bytes:
        """导出为 PDF 文档"""
        # Register CJK fonts so Chinese text renders correctly.
        # Without this, reportlab's default Type1 fonts silently drop CJK characters,
        # producing PDFs that appear empty or corrupted.
        self._register_cjk_fonts()

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72,
        )

        styles = getSampleStyleSheet()
        # Override default fonts with CJK-capable ones
        self._patch_styles_for_cjk(styles)
        story = []

        # 标题
        title_style = ParagraphStyle(
            "CustomTitle",
            parent=styles["Heading1"],
            fontSize=24,
            alignment=1,
            spaceAfter=30,
        )
        story.append(Paragraph("智能运维故障分析报告", title_style))
        story.append(Spacer(1, 20))

        # 元信息
        meta_style = ParagraphStyle("Meta", parent=styles["Normal"], fontSize=10, textColor=colors.gray)
        meta_text = (
            f"故障编号：{self.result.incident_id or 'N/A'} | "
            f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        )
        story.append(Paragraph(meta_text, meta_style))
        story.append(Spacer(1, 20))

        # 执行摘要
        story.append(Paragraph("执行摘要", styles["Heading2"]))
        summary_lines = [
            f"辩论轮次：{self.result.total_turns or 0} 轮",
            f"收敛分数：{((self.result.convergence_score or 0) * 100):.0f}%",
            f"终止原因：{self.result.termination_reason or '正常结束'}",
        ]
        for line in summary_lines:
            story.append(Paragraph(line, styles["Normal"]))

        if self.result.disputed_points:
            story.append(Paragraph("争议点：", styles["Normal"]))
            for point in self.result.disputed_points:
                story.append(Paragraph(f"• {self._escape_xml(point)}", styles["Normal"]))
            story.append(Spacer(1, 15))

        # 分析结论
        if self.result.report_text:
            story.append(Paragraph("分析结论", styles["Heading2"]))
            for line in self.result.report_text.split("\n"):
                stripped = line.strip()
                if stripped.startswith("### ") or stripped.startswith("## "):
                    heading_text = stripped.lstrip("# ")
                    story.append(Paragraph(self._escape_xml(heading_text), styles["Heading3"]))
                elif stripped.startswith("- ") or stripped.startswith("* "):
                    bullet_text = stripped[2:]
                    story.append(Paragraph(f"• {self._escape_xml(bullet_text)}", styles["Normal"]))
                elif stripped.startswith("```"):
                    continue
                elif stripped:
                    story.append(Paragraph(self._escape_xml(stripped), styles["Normal"]))
            story.append(Spacer(1, 20))

        # 辩论过程
        story.append(Paragraph("辩论过程", styles["Heading2"]))
        if self.result.history and "turns" in self.result.history:
            for i, turn in enumerate(self.result.history["turns"]):
                turn_label = f"第 {turn.get('turn_id', i + 1)} 轮 - {turn.get('round', '')}"
                agent_label = f"发言人：{turn.get('agent_id', 'N/A')}"
                turn_lines = [self._escape_xml(turn_label), self._escape_xml(agent_label)]

                if turn.get("message"):
                    content = turn["message"].get("content", "")
                    if content:
                        for cline in content[:500].split("\n"):
                            if cline.strip():
                                turn_lines.append(self._escape_xml(cline))
                        if len(content) > 500:
                            turn_lines.append("...")

                for tline in turn_lines:
                    story.append(Paragraph(tline, styles["Normal"]))
                story.append(Spacer(1, 10))

        # 页脚
        story.append(Spacer(1, 30))
        footer_style = ParagraphStyle(
            "Footer", parent=styles["Normal"], fontSize=9, textColor=colors.gray, alignment=1
        )
        story.append(Paragraph("— 由 AI-OPS 多智能体系统自动生成 —", footer_style))

        try:
            doc.build(story)
        except Exception as e:
            import logging
            logging.getLogger(__name__).error(f"PDF build failed: {e}")
            raise

        buffer.seek(0)
        return buffer.getvalue()

    @staticmethod
    def _register_cjk_fonts() -> None:
        """Register CJK-capable TrueType fonts with reportlab.

        Searches common font locations across Windows / macOS / Linux.
        Falls back gracefully if no CJK fonts are found.
        """
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        import os as _os
        import glob as _glob

        # Check if already registered
        if hasattr(ReportExporter, '_cjk_fonts_registered'):
            return

        font_search_paths = [
            # Windows
            "C:/Windows/Fonts/simsun.ttc",
            "C:/Windows/Fonts/simhei.ttf",
            "C:/Windows/Fonts/msyh.ttc",
            "C:/Windows/Fonts/msyhbd.ttc",
            # macOS
            "/System/Library/Fonts/PingFang.ttc",
            "/System/Library/Fonts/STHeiti Light.ttc",
            "/Library/Fonts/Arial Unicode.ttf",
            # Linux
            "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
            "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",
        ]

        # Also search with glob for common patterns
        _extra_globs = [
            "C:/Windows/Fonts/sim*.tt*",
            "C:/Windows/Fonts/msyh*.tt*",
            "/usr/share/fonts/**/wqy*.tt*",
            "/usr/share/fonts/**/NotoSansCJK*.tt*",
        ]
        for pattern in _extra_globs:
            try:
                for path in _glob.glob(pattern, recursive=True):
                    if path not in font_search_paths and _os.path.isfile(path):
                        font_search_paths.append(path)
            except Exception:
                pass

        for font_path in font_search_paths:
            if _os.path.isfile(font_path):
                try:
                    pdfmetrics.registerFont(TTFont("CJK", font_path))
                    ReportExporter._cjk_fonts_registered = True
                    return
                except Exception:
                    continue

        # No CJK font found — mark as attempted so we don't search again
        ReportExporter._cjk_fonts_registered = True

    @staticmethod
    def _patch_styles_for_cjk(styles) -> None:
        """Replace default Type1 font names with registered CJK font in styles.

        Only patches if a CJK font was successfully registered.
        """
        from reportlab.pdfbase import pdfmetrics
        import logging

        logger = logging.getLogger(__name__)
        try:
            pdfmetrics.getFont("CJK")
        except Exception:
            return  # No CJK font registered, keep defaults

        cjk_font_name = "CJK"
        for style_name in styles.byName:
            style = styles[style_name]
            style.fontName = cjk_font_name
        logger.info("Patched reportlab styles to use CJK font: %s", cjk_font_name)

    def save_docx(self, path: str) -> None:
        """保存为 Word 文件"""
        with open(path, 'wb') as f:
            f.write(self.export_docx())

    def save_pdf(self, path: str) -> None:
        """保存为 PDF 文件"""
        with open(path, 'wb') as f:
            f.write(self.export_pdf())


def export_debate_report(result: DebateResult, format: str = 'docx', output_path: Optional[str] = None) -> bytes:
    """
    导出辩论报告

    Args:
        result: 辩论结果
        format: 格式 ('docx' 或 'pdf')
        output_path: 可选，保存路径

    Returns:
        文件字节数据
    """
    exporter = ReportExporter(result)

    if format.lower() == 'pdf':
        data = exporter.export_pdf()
        if output_path:
            exporter.save_pdf(output_path)
    else:  # 默认 docx
        data = exporter.export_docx()
        if output_path:
            exporter.save_docx(output_path)

    return data