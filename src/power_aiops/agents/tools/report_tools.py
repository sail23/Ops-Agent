"""
报告导出工具集。

提供 Word、PDF 等格式的报告导出能力。
"""

from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any

from power_aiops.agents.tools.base import (
    Tool,
    ToolCategory,
    ToolMetadata,
    ToolResult,
)


# ─── CJK font registration for PDF export ──────────────────────────────

_cjk_font_registered = False


def _register_cjk_font() -> None:
    """Register a CJK-capable TrueType font with reportlab for Chinese text rendering."""
    global _cjk_font_registered
    if _cjk_font_registered:
        return

    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import os, glob

    font_paths = [
        "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/msyhbd.ttc",
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/Library/Fonts/Arial Unicode.ttf",
        "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
        "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",
    ]

    extra_patterns = [
        "C:/Windows/Fonts/sim*.tt*",
        "C:/Windows/Fonts/msyh*.tt*",
        "/usr/share/fonts/**/wqy*.tt*",
        "/usr/share/fonts/**/NotoSansCJK*.tt*",
    ]
    for pattern in extra_patterns:
        try:
            for path in glob.glob(pattern, recursive=True):
                if path not in font_paths and os.path.isfile(path):
                    font_paths.append(path)
        except Exception:
            pass

    for font_path in font_paths:
        if os.path.isfile(font_path):
            try:
                pdfmetrics.registerFont(TTFont("CJK", font_path))
                _cjk_font_registered = True
                return
            except Exception:
                continue

    _cjk_font_registered = True  # Mark as attempted


def _patch_styles_for_cjk(styles) -> None:
    """Override default Type1 font names with registered CJK font in styles."""
    from reportlab.pdfbase import pdfmetrics

    try:
        pdfmetrics.getFont("CJK")
    except Exception:
        return

    for style_name in styles.byName:
        style = styles[style_name]
        style.fontName = "CJK"


class MarkdownToDocxTool(Tool):
    """Markdown 转 Word (DOCX) 工具.

    将 Markdown 格式的报告转换为 Word 文档。
    纯 Python 实现，无需外部依赖。
    """

    def __init__(self) -> None:
        super().__init__()

    @property
    def metadata(self) -> ToolMetadata:
        return ToolMetadata(
            name="export_docx",
            description="将 Markdown 格式的报告内容导出为 Word (.docx) 文档。",
            category=ToolCategory.ANALYSIS,
            parameters={
                "properties": {
                    "content": {
                        "type": "string",
                        "description": "Markdown 格式的报告内容",
                    },
                    "output_path": {
                        "type": "string",
                        "description": "输出文件路径（.docx）",
                    },
                    "title": {
                        "type": "string",
                        "description": "文档标题",
                    },
                    "author": {
                        "type": "string",
                        "description": "作者",
                    },
                },
                "required": ["content", "output_path"],
            },
            examples=[
                "export_docx(content='# 故障报告\\n\\n## 概述\\n...', output_path='report.docx', title='故障分析报告')",
            ],
            tags=["export", "word", "docx", "report"],
        )

    def execute(self, **kwargs) -> ToolResult:
        content = kwargs.get("content", "")
        output_path = kwargs.get("output_path")
        title = kwargs.get("title", "故障分析报告")
        author = kwargs.get("author", "Power-AIOps")

        if not content.strip():
            return ToolResult(success=False, error="Empty content provided")

        if not output_path:
            return ToolResult(success=False, error="output_path is required")

        try:
            # 确保 .docx 扩展名
            if not output_path.lower().endswith(".docx"):
                output_path += ".docx"

            output_file = Path(output_path).expanduser().resolve()

            # 尝试使用 python-docx
            try:
                from docx import Document
                from docx.shared import Pt, Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH

                doc = Document()

                # 设置文档属性
                doc.core_properties.title = title
                doc.core_properties.author = author

                # 解析 Markdown 并转换
                self._markdown_to_docx(doc, content)

                # 保存
                doc.save(str(output_file))

                return ToolResult(
                    success=True,
                    data={
                        "path": str(output_file),
                        "format": "docx",
                        "size_bytes": output_file.stat().st_size,
                    },
                )

            except ImportError:
                # 备用方案：生成 HTML 文件
                html_content = self._markdown_to_html(content)
                html_path = str(output_file).replace(".docx", ".html")
                with open(html_path, "w", encoding="utf-8") as f:
                    f.write(html_content)

                return ToolResult(
                    success=False,
                    error=f"python-docx 未安装，已导出 HTML 格式",
                    data={
                        "path": html_path,
                        "format": "html",
                        "suggestion": "运行 pip install python-docx 以支持 DOCX 导出",
                    },
                )

        except Exception as e:
            return ToolResult(success=False, error=f"Export failed: {str(e)}")

    def _markdown_to_docx(self, doc, markdown_text: str) -> None:
        """将 Markdown 内容转换为 DOCX."""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        lines = markdown_text.split("\n")
        current_paragraph = None

        for line in lines:
            stripped = line.strip()

            # 标题
            if stripped.startswith("# "):
                p = doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("## "):
                p = doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("### "):
                p = doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("#### "):
                p = doc.add_heading(stripped[5:], level=4)
            # 代码块
            elif stripped.startswith("```"):
                current_paragraph = None  # 重置
            # 列表
            elif stripped.startswith("- ") or stripped.startswith("* "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif re.match(r"^\d+\. ", stripped):
                doc.add_paragraph(re.sub(r"^\d+\. ", "", stripped), style="List Number")
            # 分隔线
            elif stripped.startswith("---"):
                p = doc.add_paragraph()
                p.add_run("─" * 50)
            # 空行
            elif not stripped:
                current_paragraph = None
            # 普通文本
            else:
                # 处理行内格式
                text = self._process_inline_format(stripped)
                if text:
                    p = doc.add_paragraph(text)

    def _process_inline_format(self, text: str) -> str:
        """处理行内格式（粗体、斜体等）."""
        # 移除代码标记但保留内容
        text = re.sub(r"`([^`]+)`", r"\1", text)
        # 移除粗体标记但保留内容
        text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
        text = re.sub(r"__([^_]+)__", r"\1", text)
        # 移除斜体标记但保留内容
        text = re.sub(r"\*([^*]+)\*", r"\1", text)
        text = re.sub(r"_([^_]+)_", r"\1", text)
        return text

    def _markdown_to_html(self, markdown_text: str) -> str:
        """将 Markdown 转换为 HTML."""
        html_parts = [
            "<!DOCTYPE html>",
            "<html>",
            "<head>",
            "<meta charset='utf-8'>",
            "<title>故障分析报告</title>",
            "<style>",
            "body { font-family: 'Microsoft YaHei', Arial, sans-serif; max-width: 900px; margin: 0 auto; padding: 20px; }",
            "h1 { color: #333; border-bottom: 2px solid #007bff; padding-bottom: 10px; }",
            "h2 { color: #555; }",
            "code { background: #f5f5f5; padding: 2px 6px; border-radius: 3px; }",
            "pre { background: #f5f5f5; padding: 15px; border-radius: 5px; overflow-x: auto; }",
            "table { border-collapse: collapse; width: 100%; }",
            "th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }",
            "th { background-color: #007bff; color: white; }",
            "</style>",
            "</head>",
            "<body>",
        ]

        lines = markdown_text.split("\n")
        in_code_block = False

        for line in lines:
            stripped = line.strip()

            if stripped.startswith("```"):
                if in_code_block:
                    html_parts.append("</code></pre>")
                else:
                    html_parts.append("<pre><code>")
                in_code_block = not in_code_block
            elif in_code_block:
                html_parts.append(self._escape_html(line))
            elif stripped.startswith("# "):
                html_parts.append(f"<h1>{self._escape_html(stripped[2:])}</h1>")
            elif stripped.startswith("## "):
                html_parts.append(f"<h2>{self._escape_html(stripped[3:])}</h2>")
            elif stripped.startswith("### "):
                html_parts.append(f"<h3>{self._escape_html(stripped[4:])}</h3>")
            elif stripped.startswith("- ") or stripped.startswith("* "):
                html_parts.append(f"<li>{self._escape_html(stripped[2:])}</li>")
            elif re.match(r"^\d+\. ", stripped):
                html_parts.append(f"<li>{self._escape_html(re.sub(r"^\d+\. ", "", stripped))}</li>")
            elif stripped.startswith("---"):
                html_parts.append("<hr>")
            elif stripped:
                html_parts.append(f"<p>{self._escape_html(stripped)}</p>")

        if in_code_block:
            html_parts.append("</code></pre>")

        html_parts.extend(["</body>", "</html>"])
        return "\n".join(html_parts)

    def _escape_html(self, text: str) -> str:
        """转义 HTML 特殊字符."""
        return (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
        )


class MarkdownToPdfTool(Tool):
    """Markdown 转 PDF 工具.

    将 Markdown 格式的报告转换为 PDF 文档。
    """

    def __init__(self) -> None:
        super().__init__()

    @property
    def metadata(self) -> ToolMetadata:
        return ToolMetadata(
            name="export_pdf",
            description="将 Markdown 格式的报告内容导出为 PDF 文档。",
            category=ToolCategory.ANALYSIS,
            parameters={
                "properties": {
                    "content": {
                        "type": "string",
                        "description": "Markdown 格式的报告内容",
                    },
                    "output_path": {
                        "type": "string",
                        "description": "输出文件路径（.pdf）",
                    },
                    "title": {
                        "type": "string",
                        "description": "文档标题",
                    },
                },
                "required": ["content", "output_path"],
            },
            examples=[
                "export_pdf(content='# 故障报告\\n\\n## 概述\\n...', output_path='report.pdf', title='故障分析报告')",
            ],
            tags=["export", "pdf", "report"],
        )

    def execute(self, **kwargs) -> ToolResult:
        content = kwargs.get("content", "")
        output_path = kwargs.get("output_path")
        title = kwargs.get("title", "故障分析报告")

        if not content.strip():
            return ToolResult(success=False, error="Empty content provided")

        if not output_path:
            return ToolResult(success=False, error="output_path is required")

        try:
            # 确保 .pdf 扩展名
            if not output_path.lower().endswith(".pdf"):
                output_path += ".pdf"

            output_file = Path(output_path).expanduser().resolve()

            # 尝试使用 reportlab
            try:
                from reportlab.lib.pagesizes import A4
                from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
                from reportlab.lib.units import cm
                from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
                from reportlab.lib.enums import TA_LEFT
                from reportlab.pdfbase import pdfmetrics
                from reportlab.pdfbase.ttfonts import TTFont

                # Register CJK font for Chinese text rendering
                _register_cjk_font()

                # 创建文档
                doc = SimpleDocTemplate(
                    str(output_file),
                    pagesize=A4,
                    rightMargin=2 * cm,
                    leftMargin=2 * cm,
                    topMargin=2 * cm,
                    bottomMargin=2 * cm,
                )

                # 创建样式
                styles = getSampleStyleSheet()
                _patch_styles_for_cjk(styles)
                story = []

                # 添加标题
                title_style = ParagraphStyle(
                    "CustomTitle",
                    parent=styles["Heading1"],
                    fontSize=24,
                    spaceAfter=30,
                    alignment=TA_LEFT,
                )
                story.append(Paragraph(self._escape_pdf(title), title_style))
                story.append(Spacer(1, 12))

                # 解析 Markdown
                lines = content.split("\n")
                for line in lines:
                    stripped = line.strip()

                    if stripped.startswith("# "):
                        p = Paragraph(self._escape_pdf(stripped[2:]), styles["Heading1"])
                        story.append(p)
                    elif stripped.startswith("## "):
                        p = Paragraph(self._escape_pdf(stripped[3:]), styles["Heading2"])
                        story.append(p)
                    elif stripped.startswith("### "):
                        p = Paragraph(self._escape_pdf(stripped[4:]), styles["Heading3"])
                        story.append(p)
                    elif stripped.startswith("```"):
                        continue
                    elif stripped.startswith("- ") or stripped.startswith("* "):
                        p = Paragraph(self._escape_pdf(stripped[2:]), styles["Normal"])
                        story.append(p)
                    elif re.match(r"^\d+\. ", stripped):
                        p = Paragraph(self._escape_pdf(re.sub(r"^\d+\. ", "", stripped)), styles["Normal"])
                        story.append(p)
                    elif stripped.startswith("---"):
                        story.append(Spacer(1, 10))
                    elif stripped:
                        p = Paragraph(self._escape_pdf(stripped), styles["Normal"])
                        story.append(p)

                # 构建 PDF
                doc.build(story)

                return ToolResult(
                    success=True,
                    data={
                        "path": str(output_file),
                        "format": "pdf",
                        "size_bytes": output_file.stat().st_size,
                    },
                )

            except ImportError:
                return ToolResult(
                    success=False,
                    error="reportlab 未安装",
                    data={
                        "suggestion": "运行 pip install reportlab 以支持 PDF 导出",
                        "alternative": "export_docx 工具可生成 Word 文档",
                    },
                )

        except Exception as e:
            return ToolResult(success=False, error=f"PDF export failed: {str(e)}")

    def _apply_inline_formatting(self, text: str) -> str:
        """XML-escape first, then convert Markdown inline formatting to reportlab XML tags.

        Escaping before conversion ensures that stray < > & in the original text
        don't get mistaken for XML tags. Only **bold**, *italic*, and `code`
        are converted — underscore-based formatting is intentionally excluded
        because `_` is common in technical identifiers.
        """
        from xml.sax.saxutils import escape as _xml_escape

        text = _xml_escape(text)

        # Code: `text` → <font face="Courier">text</font>
        text = re.sub(r"`([^`]+)`", r"<font face='Courier'>\1</font>", text)
        # Bold: **text** → <b>text</b>
        text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
        # Italic: *text* → <i>text</i>
        text = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"<i>\1</i>", text)

        return text

    def _escape_pdf(self, text: str) -> str:
        """转义 PDF 特殊字符（仅文本内容，不处理 XML 标签）."""
        return self._apply_inline_formatting(text)


class ReportTemplateTool(Tool):
    """报告模板工具.

    提供故障报告的标准模板。
    """

    def __init__(self) -> None:
        super().__init__()

    @property
    def metadata(self) -> ToolMetadata:
        return ToolMetadata(
            name="get_report_template",
            description="获取故障分析报告的标准模板（Markdown 格式）。",
            category=ToolCategory.ANALYSIS,
            parameters={
                "properties": {
                    "template_type": {
                        "type": "string",
                        "description": "模板类型: incident（故障报告）, postmortem（复盘报告）, daily（日报）",
                    },
                },
                "required": ["template_type"],
            },
            examples=[
                'get_report_template(template_type="incident")',
                'get_report_template(template_type="postmortem")',
            ],
            tags=["template", "report", "incident"],
        )

    def execute(self, **kwargs) -> ToolResult:
        template_type = kwargs.get("template_type", "incident")

        templates = {
            "incident": self._incident_template(),
            "postmortem": self._postmortem_template(),
            "daily": self._daily_template(),
        }

        if template_type not in templates:
            return ToolResult(
                success=False,
                error=f"Unknown template type: {template_type}. Available: {list(templates.keys())}",
            )

        return ToolResult(
            success=True,
            data={
                "type": template_type,
                "template": templates[template_type],
            },
        )

    def _incident_template(self) -> str:
        return """# 故障分析报告

## 基本信息
- **故障编号**: {incident_id}
- **故障时间**: {incident_time}
- **影响范围**: {impact_scope}
- **严重程度**: {severity}

## 故障摘要
{summary}

## 时间线
| 时间 | 事件 | 负责人 |
|------|------|--------|
| {timeline} | | |

## 影响分析
### 业务影响
{business_impact}

### 技术影响
{technical_impact}

## 根因分析
{root_cause}

## 处置过程
{resolution}

## 改进措施
### 短期措施
1.

### 长期措施
1.

## 附录
### 相关日志
```

### 相关配置变更
```
"""

    def _postmortem_template(self) -> str:
        return """# 故障复盘报告

## 概述
{overview}

## 故障统计
- 故障持续时间: {duration}
- 影响用户数: {affected_users}
- 损失估算: {loss_estimate}

## 故障时间线
```
{chronology}
```

## 根因分析（5 Why）
**Why 1**: {why1}
**Why 2**: {why2}
**Why 3**: {why3}
**Why 4**: {why4}
**Why 5**: {root_cause}

## 处置回顾
### 做得好的地方
- {good_practices}

### 需要改进的地方
- {improvements}

## 行动项
| 行动项 | 负责人 | 完成日期 | 状态 |
|--------|--------|----------|------|
| | | | |

## 经验教训
{lessons_learned}
"""

    def _daily_template(self) -> str:
        return """# 运维日报 - {date}

## 今日概览
| 指标 | 数值 |
|------|------|
| 告警数 | |
| 故障数 | |
| 变更数 | |

## 重要事件
### 故障处理
- {incidents}

### 变更记录
- {changes}

## 监控状态
{monitoring}

## 明日计划
- [ ]

## 备注
{notes}
"""
